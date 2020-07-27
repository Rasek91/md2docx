from argparse import ArgumentParser
from docx import Document, opc, oxml
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor
from glob import glob
from html.parser import HTMLParser
from io import BytesIO
from mistune import markdown
from pygments import highlight
from pygments.formatters import HtmlFormatter
from pygments.lexers import get_lexer_by_name
from requests import get
from os import path
from yaml import safe_load

class MDConverter(HTMLParser):
    """
    Create docx file from a list of MD files.
    
    Attributes:
      - config: the name of the YAML config file
    """
    
    #init function
    def __init__(self, config):
        #load the config file
        self.parse_config(config)
        #variables to save the tags during handle_data
        self.highlighting = False
        self.paragraph = None
        self.style = None
        self.last = None
        self.last_list_style = None
        self.link = None
        self.table = None
        self.table_row = 0
        self.table_column = 0
        self.list_level = 0
        #list of tags to handle
        self.type_tags = ["h1", "h2", "h3", "h4", "h5", "h6", "p", "pre", "table"]
        self.list_tags = ["ol", "ul"]
        self.style_tags = ["code", "del", "em", "strong"]
        self.table_tags = ["th", "td", "tr"]
        self.link_tags = ["a"]
        #load the template document
        self.document = Document(self.config["Template"])
        #create a codehighlighter instance
        self.codehighlighter = CodeHighlighter(self.document)
        #call the init function of the super
        super().__init__()
    
    #function to parse config file
    def parse_config(self, config):
        with open(config, "r") as stream:
            #safe load the YAML config file content
            self.config = safe_load(stream)
    
    #function to do the parsing
    def do(self):
        self.parse_mds()
        self.generate()
    
    #function to parse MD files
    def parse_mds(self):
        #create a sorted list of MD files
        mds = glob(path.join(self.config["Folder"], "*.md"))
        mds.sort()
        
        #parse all the files
        for filename in mds:
            self.parse_md(filename)
    
    #function to parse one MD file
    def parse_md(self, filename):
        #create HTML from MD file
        parsed = markdown(open(filename).read())
        #feed the HTML version to the HTMLParser
        self.feed(parsed)
    
    #function to add paragraph
    def create_paragraph(self, tag):
        #create haeding paragraph
        if 'h' in self.paragraph:
            self.last = self.document.add_heading("", level = int(self.paragraph[1]))
        #create a new paragraph
        elif self.paragraph == 'p':
            self.last = self.document.add_paragraph("")
        #create a code paragraph
        elif self.paragraph == 'pre':
            self.last = self.document.add_paragraph("", style = "code")
        #create unordered list paragraph
        elif self.paragraph == 'ul':
            if tag == 'ul':
                self.list_level += 1
            elif tag == 'li':
                self.last = self.document.add_paragraph("", style = "unordered list{}".format(self.list_level))
        #create unordered list paragraph
        elif self.paragraph == 'ol':
            if tag == 'ol':
                self.list_level += 1
            elif tag == 'li':
                self.last = self.document.add_paragraph("", style = "ordered list{}".format(self.list_level))
        #create a new table
        elif self.paragraph == 'table':
            self.table = self.document.add_table(rows = 1, cols = 1, style = self.config["Table Style"])
            
            for cell in self.table.columns[0].cells:
                cell.width = Cm(2.0)
            
            self.table.autofit = True
    
    #function to add hyperlink
    #https://stackoverflow.com/a/47666747
    def add_hyperlink(self, text, url):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = self.last.part
        r_id = part.relate_to(url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external = True)
        # Create the w:hyperlink tag and add needed values
        hyperlink = oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(oxml.shared.qn("r:id"), r_id, )
        # Create a w:r element and a new w:rPr element
        new_run = oxml.shared.OxmlElement("w:r")
        rPr = oxml.shared.OxmlElement("w:rPr")
        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        # Create a new Run object and add the hyperlink into it
        r = self.last.add_run()
        r._r.append(hyperlink)
        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True
    
    #function to add horizontal ruler
    #https://github.com/python-openxml/python-docx/issues/105#issuecomment-442786431
    def inserthr(self):
        p = self.document.paragraphs[-1]._p  # p is the <w:p> XML element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        pPr.insert_element_before(pBdr,
            "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
            "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
            "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
            "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
            "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
            "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr",
            "w:pPrChange"
        )
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
    
    #function to add style to the last run
    def add_style(self):
        #get the font of the last run
        font = self.last.runs[-1].font
        
        #bold
        if self.style == 'strong':
            font.bold = True
        #italic
        elif self.style == 'em':
            font.italic = True
        #strike
        elif self.style == 'del':
            font.strike = True
        #inline code
        elif self.paragraph != 'pre' and self.style == 'code':
            font.name = self.config["Inline Code Font"]
            font.color.rgb = RGBColor(int(self.config["Inline Code Color"][0:2], 16), int(self.config["Inline Code Color"][2:4], 16), int(self.config["Inline Code Color"][4:], 16))
    
    #function to add image to the document
    def add_image(self, image):
        #download the image if it is a link
        if image.startswith("http://") or image.startswith("https://"):
            response = get(image)
            image = BytesIO(response.content)
        
        self.document.add_picture(image, width = Cm(8.0))
    
    #function to handle start of an HTML tag
    def handle_starttag(self, tag, attrs):
        #add a new paragraph
        if (tag in self.type_tags or tag in self.list_tags) and self.highlighting == False:
            if self.paragraph == None:
                self.paragraph = tag
            
            self.create_paragraph(tag)
        #add code syntax highlight
        elif tag == 'code' and self.paragraph == 'pre' and len(attrs) != 0 and attrs[0][0] == 'class':
            self.highlighting = [attrs[0][1].split("-")[-1], self.config["Syntax Highlight Style"]]
        #add new line to a list
        elif tag == 'li':
            self.create_paragraph(tag)
        #save style tag
        elif tag in self.style_tags:
            self.style = tag
        #save link tag
        elif tag in self.link_tags:
            self.link = [tag, next(link[1] for link in attrs if link[0] == "href")]
        #add horizontal ruler
        elif tag == 'hr':
            self.inserthr()
        #add image
        elif tag == 'img':
            self.add_image(next(link[1] for link in attrs if link[0] == "src"))
        #save table tag
        elif tag in self.table_tags:
            self.paragraph = tag
            
            if tag == 'tr' and self.table_row != 0:
                self.table.add_row()
    
    #function to handle end of an HTML tag
    def handle_endtag(self, tag):
        #reset list tags
        if tag in self.list_tags:
            self.list_level -= 1
            
            if tag == 'ol':
                self.last_list_style = None
            
            if self.list_level == 0:
                self.paragraph = None
                self.last = None
        #reset code syntax highlighter
        elif tag == 'code' and self.paragraph == 'pre' and self.highlighting != False:
            self.highlighting = False
        #reset paragraph tags
        elif tag in self.type_tags and tag == self.paragraph:
            self.paragraph = None
            self.last = None
        #reset table
        elif tag == 'table':
            #set table lenght to not be full page lenght
            #https://github.com/python-openxml/python-docx/issues/315#issuecomment-239259678
            for r in self.table.rows:
                for c in r._tr.tc_lst:
                    tcW = c.tcPr.tcW
                    tcW.type = 'auto'
                    tcW.w = 0
            
            self.table = None
            self.table_row = 0
        #reset style tags
        elif tag in self.style_tags:
            self.style = None
        #reset link tags
        elif tag in self.link_tags:
            self.link = None
        #reset table tags
        elif tag in self.table_tags:
            self.paragraph = None
            
            if tag == 'tr':
                self.table_column = 0
                self.table_row += 1
    
    #function to handle data inside an HTML tag
    def handle_data(self, data):
        #add text to not a table
        if self.table == None:
            #add text if it not a new line only
            if data != '\n' and data:
                #do code syntax highlighting
                if self.highlighting != False:
                    self.codehighlighter.feed(highlight(data, get_lexer_by_name(self.highlighting[0]), HtmlFormatter(style = self.highlighting[1], noclasses = True)))
                #add text to the last paragraph
                elif self.last != None:
                    #add simple text
                    if self.link == None:
                        self.last.add_run(data)
                    #add link
                    elif self.link[0] == 'a':
                        self.add_hyperlink(data, self.link[1])
                
                #add style to the last text
                if self.style != None:
                    self.add_style()
        #add text to a table
        else:
            #merge cell to the upper one
            if data == '^^':
                self.table.rows[self.table_row].cells[self.table_column].merge(self.table.rows[self.table_row - 1].cells[self.table_column])
                self.table_column += 1
            #merge cell to the left one
            elif data == '<<':
                self.table.rows[self.table_row].cells[self.table_column].merge(self.table.rows[self.table_row].cells[self.table_column - 1])
                self.table_column += 1
            #add text to the header
            elif self.paragraph == 'th':
                #add new column if needed
                if self.table_column != 0:
                    self.table.add_column(width = Cm(2.0))
                
                self.table.rows[self.table_row].cells[self.table_column].text = data
                self.table_column += 1
            #add text to the table body
            elif self.paragraph == 'td':
                self.table.rows[self.table_row].cells[self.table_column].text = data
                self.table_column += 1
    
    #function to generate the docx file from the tamplate
    def generate(self):
        #call the function of Document to save the new file
        self.document.save(self.config["Filename"])
    
    #delete function
    def __del__(self):
        pass

class CodeHighlighter(HTMLParser):
    """
    Add code syntax highlighting to a code paragraph.
    
    Attributes:
      - document: document where the highlighting needed
    """
    
    #init function
    def __init__(self, document):
        self.document = document
        #call the init function of the super
        super().__init__()
    
    #function to handle start of an HTML tag
    def handle_starttag(self, tag, attrs):
        #only span tags are in the generated document
        if tag == 'span':
            #if there is highlighting change the font
            if len(attrs) != 0:
                styles = next(attribute[1] for attribute in attrs if attribute[0] == "style")
                font = self.document.paragraphs[-1].add_run().font
                
                for style in styles.split(";"):
                    if style.split(": ")[0] == 'color':
                        color = style.split(": ")[1][1:]
                        font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:], 16))
                    elif style.split(": ")[0] == 'font-weight' and style.split(": ")[1] == 'bold':
                        font.bold = True
            #if there is no highlighting add an unchanged font
            else:
                self.document.paragraphs[-1].add_run()
    
    #function to handle end of an HTML tag
    def handle_endtag(self, tag):
        pass
    
    #function to handle data inside an HTML tag
    def handle_data(self, data):
        for character in data:
            #add break if needed
            if character == '\n':
                self.document.paragraphs[-1].runs[-1].add_break()
            #add the text to the last run
            else:
                self.document.paragraphs[-1].runs[-1].add_text(character)
    
    #delete function
    def __del__(self):
        pass

#function to parse the arguments
def parse_arguments():
    #create instance of ArgumentParser class
    parser = ArgumentParser(description = "MD 2 DOCX Converter")
    #reset the groups
    parser._action_groups.pop()
    #add new required group
    required = parser.add_argument_group("Required Arguments")
    #add the config parser
    required.add_argument("-c", "--config", type = str, required = True, help = "The configuration file of MD 2 DOCX Converter")
    #parse arguments
    arguments = parser.parse_args()
    
    #return with the config file name
    return arguments.config

#main function
def main():
    #parse configuration
    config = parse_arguments()
    #create instance of MDConverter class
    md2docx = MDConverter(config)
    #do the converting
    md2docx.do()

#run main if the script call directly
if __name__ == '__main__':
    main()

